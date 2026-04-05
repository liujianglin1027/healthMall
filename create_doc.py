from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('如何通过Trae克隆Git仓库并运行应用教程', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('目录', level=1)
doc.add_paragraph('1. 什么是Trae')
doc.add_paragraph('2. 安装和启动Trae')
doc.add_paragraph('3. 通过Trae克隆Git仓库')
doc.add_paragraph('4. 在Trae中运行应用')
doc.add_paragraph('5. 常见问题和解决方案')

doc.add_heading('1. 什么是Trae', level=1)
doc.add_paragraph('Trae是一款由国内团队开发的AI驱动的集成开发环境(IDE)，它结合了传统IDE的功能和AI辅助编程能力，旨在帮助开发者更高效地进行代码开发、项目管理和问题解决。')
doc.add_paragraph()
doc.add_paragraph('Trae的主要特点包括：')
doc.add_paragraph('• 智能代码补全和建议', style='List Bullet')
doc.add_paragraph('• 代码解释和文档生成', style='List Bullet')
doc.add_paragraph('• 项目管理和版本控制集成', style='List Bullet')
doc.add_paragraph('• 终端和命令行工具', style='List Bullet')
doc.add_paragraph('• 多语言支持', style='List Bullet')

doc.add_heading('2. 安装和启动Trae', level=1)
doc.add_heading('安装Trae', level=2)
doc.add_paragraph('1. 访问Trae官方网站 (https://trae.cn/)')
doc.add_paragraph('2. 下载适合您操作系统的安装包')
doc.add_paragraph('3. 运行安装程序，按照提示完成安装过程')
doc.add_paragraph('4. 安装完成后，启动Trae应用')

doc.add_heading('首次启动Trae', level=2)
doc.add_paragraph('1. 打开Trae应用')
doc.add_paragraph('2. 登录您的Trae账户（如果没有账户，按照提示注册一个）')
doc.add_paragraph('3. 进入Trae的主界面')

doc.add_heading('3. 通过Trae克隆Git仓库', level=1)
doc.add_heading('方法一：通过Trae界面克隆', level=2)
doc.add_paragraph('1. 在Trae主界面，点击左侧导航栏中的「项目」')
doc.add_paragraph('2. 点击「克隆项目」按钮')
doc.add_paragraph('3. 在弹出的对话框中，输入Git仓库的URL')
doc.add_paragraph('4. 选择本地保存路径')
doc.add_paragraph('5. 点击「克隆」按钮，等待克隆完成')

doc.add_heading('方法二：通过终端克隆', level=2)
doc.add_paragraph('1. 在Trae主界面，点击下方的「终端」按钮，打开集成终端')
doc.add_paragraph('2. 在终端中，使用Git命令克隆仓库：')
doc.add_paragraph('git clone <仓库URL> <本地路径>')
doc.add_paragraph('3. 等待克隆完成')

doc.add_heading('克隆后打开项目', level=2)
doc.add_paragraph('1. 克隆完成后，Trae会自动提示您打开项目')
doc.add_paragraph('2. 点击「打开项目」按钮，或在左侧导航栏中找到克隆的项目并点击打开')

doc.add_heading('4. 在Trae中运行应用', level=1)
doc.add_heading('运行Python应用', level=2)
doc.add_paragraph('1. 打开克隆的Python项目')
doc.add_paragraph('2. 确保项目中包含必要的依赖文件（如requirements.txt）')
doc.add_paragraph('3. 在终端中安装依赖：pip install -r requirements.txt')
doc.add_paragraph('4. 找到项目的主入口文件（通常是main.py或app.py）')
doc.add_paragraph('5. 点击文件编辑器右上角的「运行」按钮，或在终端中运行：python <主入口文件>')

doc.add_heading('运行前端应用（如React、Vue等）', level=2)
doc.add_paragraph('1. 打开克隆的前端项目')
doc.add_paragraph('2. 确保项目中包含package.json文件')
doc.add_paragraph('3. 在终端中安装依赖：npm install')
doc.add_paragraph('4. 运行开发服务器：npm run dev')
doc.add_paragraph('5. 打开浏览器访问终端中显示的本地服务器地址')

doc.add_heading('运行其他类型的应用', level=2)
doc.add_paragraph('对于其他类型的应用（如Node.js、Java等），请参考项目的README文件或文档，按照相应的步骤运行应用。')

doc.add_heading('5. 常见问题和解决方案', level=1)
doc.add_heading('问题1：克隆仓库失败', level=2)
doc.add_paragraph('解决方案：')
doc.add_paragraph('• 检查网络连接是否正常', style='List Bullet')
doc.add_paragraph('• 确认Git仓库URL是否正确', style='List Bullet')
doc.add_paragraph('• 检查是否有足够的权限访问该仓库', style='List Bullet')
doc.add_paragraph('• 尝试使用HTTPS而不是SSH协议克隆', style='List Bullet')

doc.add_heading('问题2：依赖安装失败', level=2)
doc.add_paragraph('解决方案：')
doc.add_paragraph('• 检查网络连接是否正常', style='List Bullet')
doc.add_paragraph('• 尝试使用国内镜像源（如pip使用豆瓣源，npm使用淘宝源）', style='List Bullet')
doc.add_paragraph('• 确认依赖版本是否与当前环境兼容', style='List Bullet')

doc.add_heading('问题3：应用运行失败', level=2)
doc.add_paragraph('解决方案：')
doc.add_paragraph('• 检查应用的错误信息，定位问题原因', style='List Bullet')
doc.add_paragraph('• 确认所有依赖都已正确安装', style='List Bullet')
doc.add_paragraph('• 检查环境变量是否配置正确', style='List Bullet')
doc.add_paragraph('• 参考项目的README文件或文档，确认运行步骤是否正确', style='List Bullet')

doc.add_heading('问题4：Trae响应缓慢', level=2)
doc.add_paragraph('解决方案：')
doc.add_paragraph('• 检查系统资源使用情况，确保有足够的内存和CPU资源', style='List Bullet')
doc.add_paragraph('• 关闭不必要的应用程序', style='List Bullet')
doc.add_paragraph('• 尝试重启Trae应用', style='List Bullet')
doc.add_paragraph('• 如果问题持续存在，联系Trae官方支持', style='List Bullet')

doc.add_heading('总结', level=1)
doc.add_paragraph('通过本教程，您应该已经了解了如何通过Trae克隆Git仓库并运行应用的基本步骤。Trae作为一款AI驱动的IDE，不仅提供了传统IDE的功能，还能通过AI辅助提高开发效率。')
doc.add_paragraph()
doc.add_paragraph('如果您在使用过程中遇到任何问题，可以参考本教程的常见问题部分，或访问Trae官方文档和社区寻求帮助。')
doc.add_paragraph()
doc.add_paragraph('祝您使用Trae开发愉快！')

doc.save(r'C:\Users\dingmengmeng\Desktop\如何通过Trae克隆Git仓库并运行应用教程.docx')
print('文档已成功保存到桌面！')
